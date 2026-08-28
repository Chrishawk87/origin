"""gaps.py — Origin Gap Finder (INTERNAL tool).

Load a contractor's documents plus their industry / state (and, optionally, the
operators they work for) and get back EVERY compliance gap:

  * MISSING  — a required standard nothing in their docs covers at all
  * FAILING  — a required standard they have SOMETHING for, but it's missing
               required elements a reviewer will bounce
  * PRESENT  — a required standard their docs already satisfy

Each gap carries the specific missing elements, the reviewer failure points that
will sink it, and whether Origin can auto-draft the fix (Phase 2).

This is the orchestration layer over machinery that ALREADY exists in Origin:
  - compliance_kb.naics_applicable()  -> required standard set by industry/state
  - compliance_kb.hiring_client_gaps()-> operator-specific extra requirements
  - document_tools.extract_text()     -> read each file (PDF/scan/docx/img, OCR)
  - compliance_kb.resolve_standards() -> which standards a document invokes
  - compliance_kb.check_standard()    -> element-by-element coverage of a standard
  - compliance_kb.render_program()    -> the fillable template (Phase 2 auto-draft)

Nothing here is client-facing. It runs behind the app access token.
"""

from __future__ import annotations

import re
from typing import Any, Dict, List, Optional, Tuple

from . import compliance_kb as kb

try:
    from .tools import document_tools as _doc
except Exception:  # pragma: no cover - tools package layout guard
    _doc = None

# A standard counts as PRESENT once its required elements are this well covered.
# Mirrors validate_document()'s default gate so the two agree.
PASS_RATIO = 0.8

# check_standard() is keyword-based, so generic safety words ("training",
# "program", "employee") bleed a little coverage into standards the contractor
# never actually wrote. To avoid mislabelling those as FAILING (they're really
# MISSING), a standard is only FAILING if the docs EITHER explicitly invoke it
# (matched by citation/title via resolve_standards) OR clear this coverage floor.
FAILING_FLOOR = 0.34

# Prequal platform fingerprints — same idea compliance_intake uses, so the
# report can say "these look like ISNetworld documents" without being asked.
_PLATFORM_HINTS = {
    "ISNetworld": ["isnetworld", "isn ", "isnet", " ravs", "review and verification"],
    "Avetta": ["avetta", "browz"],
    "Veriforce": ["veriforce", "pec safety", "pec premier"],
    "PEC": ["pec safety", "pec premier"],
    "ComplyWorks": ["complyworks"],
}


def extract_text(path: str) -> str:
    """Read one uploaded file to text (PDF/scan/docx/image/plain)."""
    if _doc is None:
        return f"ERROR: document reader unavailable (cannot read {path})"
    return _doc.extract_text(path)


def _detect_platforms(text: str) -> List[str]:
    low = text.lower()
    found = []
    for name, hints in _PLATFORM_HINTS.items():
        if any(h in low for h in hints):
            found.append(name)
    return found


def _severity(status: str, needs_program: bool, category: str) -> int:
    """Lower sorts first. Missing written programs are the most urgent."""
    if status == "MISSING":
        return 0 if needs_program else 1
    if status == "FAILING":
        return 2 if needs_program else 3
    return 5  # PRESENT


# ── Phase 3: parse an uploaded ISN / Avetta deficiency report ───────────────
# The contractor exports their platform's deficiency / RAVS review / scorecard
# and drops it in. These lines carry a status the platform assigned — when it's
# one of these "flagged" states, that requirement is a confirmed problem.
_DEFICIENCY_NEG = (
    "not met", "not-met", "deficient", "deficiency", "expired", "rejected",
    "reject", "missing", "failed", "incomplete", "not submitted", "not provided",
    "under review", "needs revision", "non-compliant", "noncompliant", "past due",
    "overdue", "not accepted", "revise", "not on file", "no document",
)
_DEFICIENCY_POS = (
    "met", "approved", "accepted", "compliant", "complete", "current", "active",
    "verified", "satisfactory", "passed", "up to date", "up-to-date",
)
_REPORT_SIGNALS = (
    "deficienc", "not met", "ravs", "requirement", "verification", "under review",
    "expired", "rejected", "scorecard", "action item", "non-compliant", "review status",
)


def looks_like_deficiency_report(text: str) -> bool:
    """True when a document reads like a platform's deficiency/review export
    rather than a written program the contractor authored."""
    low = (text or "").lower()
    if not low:
        return False
    platform = any(h in low for hints in _PLATFORM_HINTS.values() for h in hints)
    neg_hits = sum(low.count(k) for k in
                   ("not met", "deficien", "expired", "rejected", "under review",
                    "incomplete", "missing", "not submitted"))
    signals = sum(1 for s in _REPORT_SIGNALS if s in low)
    return (platform and neg_hits >= 2) or (signals >= 3 and neg_hits >= 2)


def _clean_topic(line: str) -> str:
    t = line
    for k in _DEFICIENCY_NEG + _DEFICIENCY_POS:
        t = re.sub(r"(?i)\b" + re.escape(k) + r"\b", " ", t)
    t = re.sub(r"[|\t;]+", " ", t)
    t = re.sub(r"\s{2,}", " ", t)
    return t.strip(" -:•\t.")


# Words too generic to identify a standard by (they appear in dozens of titles).
_TOPIC_STOP = {
    "program", "plan", "policy", "procedure", "procedures", "written", "the",
    "of", "and", "for", "a", "an", "to", "in", "on", "safety", "management",
    "control", "protection", "general", "industry", "requirement", "requirements",
    "compliance", "standard", "review", "some", "custom", "client",
}


def _topic_tokens(s: str) -> set:
    return {w for w in re.findall(r"[a-z0-9]+", (s or "").lower())
            if len(w) > 2 and w not in _TOPIC_STOP}


def _match_standard(topic: str) -> Optional[dict]:
    """Map a free-text deficiency line to a KB standard.

    Strong signal first (resolve_standards, which matches by citation/title).
    Otherwise pull the top keyword-search candidates and RE-RANK them by how
    well the standard's title overlaps the topic words — plain keyword search
    over-weights long titles (e.g. 'Hazard Communication' wrongly matching
    'Hazardous Waste Operations…' on the shared 'hazard' stem), so a direct
    title-token comparison is a better tiebreaker."""
    if len(topic) < 4:
        return None
    hits = kb.resolve_standards(topic, limit=1)
    if hits:
        return hits[0]
    res = kb.search(topic, limit=6)
    if not res:
        return None
    twords = _topic_tokens(topic)
    if not twords:
        return res[0]
    best, best_score = res[0], -1.0
    for i, r in enumerate(res):
        cand = _topic_tokens(r.get("title", ""))
        if not cand:
            continue
        overlap = len(twords & cand)
        # overlap count dominates; normalize by topic size; nudge by search rank.
        score = overlap + (len(twords & cand) / len(twords)) - i * 0.01
        if score > best_score:
            best, best_score = r, score
    # If nothing shares a meaningful word with the topic, don't force a match.
    if best_score <= 0:
        return None
    return best


def parse_deficiency_report(text: str) -> List[dict]:
    """Pull the flagged line-items out of a deficiency report and map each to a
    KB standard where possible."""
    items: List[dict] = []
    seen: set = set()
    for raw in (text or "").splitlines():
        line = raw.strip()
        if len(line) < 4:
            continue
        low = line.lower()
        neg = next((k for k in _DEFICIENCY_NEG if k in low), None)
        if not neg:
            continue
        topic = _clean_topic(line)
        if len(topic) < 4:
            continue
        std = _match_standard(topic)
        key = (std["id"] if std else topic.lower())
        if key in seen:
            continue
        seen.add(key)
        items.append({
            "raw": line[:300],
            "status": neg,
            "topic": topic[:160],
            "standard_id": std["id"] if std else None,
            "matched_title": std.get("title") if std else None,
            "citation": (std.get("citation", "") if std else ""),
        })
    return items


def _required_set(industry: str, state: Optional[str],
                  operators: Optional[List[str]]) -> Tuple[List[dict], dict]:
    """Union the NAICS/state baseline with any operator overlays.

    Returns (list of standard stubs {id,title,citation,category,written_program},
    meta) where meta records the sector, state, operators matched, and notes.
    """
    base = kb.naics_applicable(industry, state=state)
    stubs: Dict[str, dict] = {}
    for s in base.get("standards", []):
        stubs[s["id"]] = s

    op_meta: List[dict] = []
    for name in (operators or []):
        name = (name or "").strip()
        if not name:
            continue
        og = kb.hiring_client_gaps(name, industry, state=state)
        if og.get("error"):
            op_meta.append({"operator": name, "matched": False, "note": og["error"]})
            continue
        for s in og.get("baseline", {}).get("standards", []):
            stubs.setdefault(s["id"], s)
        op_meta.append({
            "operator": og.get("hiring_client", name),
            "matched": True,
            "confirmed": og.get("confirmed", False),
            "overlay": og.get("overlay", {}),
            "note": og.get("note", ""),
        })

    meta = {
        "industry": industry,
        "sector": base.get("sector"),
        "sector_label": base.get("sector_label"),
        "state": base.get("state"),
        "gap_note": base.get("gap_note"),
        "operators": op_meta,
        "baseline_count": base.get("count", 0),
    }
    return list(stubs.values()), meta


def _needs_program(stub: dict) -> bool:
    return (stub.get("written_program", "") or "").strip().lower() in ("yes", "conditional")


def find_gaps(
    industry: str,
    state: Optional[str] = None,
    operators: Optional[List[str]] = None,
    docs: Optional[List[Dict[str, str]]] = None,
) -> dict:
    """Compute the full gap report.

    ``docs`` is a list of ``{"name": <filename>, "text": <extracted text>}``.
    (The route extracts the text; the engine stays pure so it's easy to test.)
    """
    docs = docs or []

    # Phase 3: split the uploaded docs into two piles. A platform deficiency
    # report (ISN/Avetta export, RAVS review, scorecard) is NOT evidence the
    # contractor covers a standard — it's a list of what the reviewer already
    # rejected. Counting its text as "coverage" would falsely raise the score
    # (the report literally names the programs it's flagging). So the coverage
    # corpus is built from the CONTRACTOR'S OWN documents only; reports feed the
    # flagged-items overlay instead.
    program_docs: List[Dict[str, str]] = []
    report_docs: List[Dict[str, str]] = []
    for d in docs:
        if looks_like_deficiency_report(d.get("text", "")):
            report_docs.append(d)
        else:
            program_docs.append(d)

    corpus = "\n\n".join(d.get("text", "") for d in program_docs)

    # Parse every flagged line-item out of the report(s) and index by standard.
    flagged_items: List[dict] = []
    for d in report_docs:
        for item in parse_deficiency_report(d.get("text", "")):
            item = dict(item)
            item["source"] = d.get("name", "")
            flagged_items.append(item)
    # standard_id -> the flagged item (first wins; keep the reason for display).
    flagged_by_id: Dict[str, dict] = {}
    for it in flagged_items:
        sid = it.get("standard_id")
        if sid and sid not in flagged_by_id:
            flagged_by_id[sid] = it

    required, meta = _required_set(industry, state, operators)

    # A platform can flag a standard that the trade baseline didn't surface —
    # union those in so a confirmed deficiency is never silently dropped.
    required_ids_have = {s["id"] for s in required}
    for sid in flagged_by_id:
        if sid in required_ids_have:
            continue
        rec = kb.get(sid)
        if not rec:
            continue
        required.append({
            "id": sid,
            "title": rec.get("title", ""),
            "citation": rec.get("citation", ""),
            "category": rec.get("category", ""),
            "written_program": rec.get("written_program", ""),
        })

    # Which standards do the contractor's OWN docs actually invoke? (informational —
    # surfaces things they cover that we didn't flag as required for the trade.)
    resolved_ids = {r["id"] for r in kb.resolve_standards(corpus, limit=50)} if corpus else set()

    gaps: List[dict] = []
    counts = {"MISSING": 0, "FAILING": 0, "PRESENT": 0}

    for stub in required:
        rec = kb.get(stub["id"])
        if not rec:
            continue
        needs_prog = _needs_program(stub)
        chk = kb.check_standard(rec, corpus) if corpus else {
            "elements_total": len(rec.get("required_elements", [])),
            "elements_covered": 0,
            "coverage_ratio": 0.0,
            "missing_elements": list(rec.get("required_elements", [])),
            "failure_points": rec.get("failure_points", []),
            "training": rec.get("training", ""),
            "recordkeeping": rec.get("recordkeeping", ""),
        }
        ratio = chk.get("coverage_ratio", 0.0)
        invoked = stub["id"] in resolved_ids  # cited/titled — a strong signal

        if ratio >= PASS_RATIO:
            status = "PRESENT"
        elif invoked or ratio >= FAILING_FLOOR:
            status = "FAILING"
        else:
            status = "MISSING"

        # Phase 3 override: the platform already ruled on this standard. Trust it
        # over our text-coverage guess — never show PRESENT for something the
        # reviewer rejected. If our docs cover it (PRESENT/FAILING) it's FAILING
        # (they have a draft but the platform bounced it); if we have nothing
        # it stays MISSING. Either way it jumps to the top of the queue.
        flag = flagged_by_id.get(stub["id"])
        if flag:
            status = "MISSING" if status == "MISSING" else "FAILING"
        counts[status] += 1

        gap = {
            "id": stub["id"],
            "title": stub.get("title", rec.get("title", "")),
            "citation": stub.get("citation", rec.get("citation", "")),
            "category": stub.get("category", rec.get("category", "")),
            "written_program": stub.get("written_program", ""),
            "needs_program": needs_prog,
            "status": status,
            "coverage_ratio": ratio,
            "elements_total": chk.get("elements_total", 0),
            "elements_covered": chk.get("elements_covered", 0),
            "missing_elements": chk.get("missing_elements", []),
            "failure_points": chk.get("failure_points", []),
            "training": chk.get("training", ""),
            "recordkeeping": chk.get("recordkeeping", ""),
            # Phase 2: Origin can auto-draft any written-program standard from KB.
            "can_autodraft": needs_prog,
            "platform_flagged": bool(flag),
            "platform_reason": (flag.get("raw") if flag else ""),
            "platform_status": (flag.get("status") if flag else ""),
            "platform_source": (flag.get("source") if flag else ""),
            "severity": _severity(status, needs_prog, stub.get("category", "")),
        }
        # A confirmed platform deficiency outranks an inferred one of the same
        # status — pull it to the very front of its severity tier.
        if flag:
            gap["severity"] -= 0.5
        gaps.append(gap)

    gaps.sort(key=lambda g: (g["severity"], g["category"], g["title"]))

    # Standards the docs cover that aren't in the required set (context only).
    extra = []
    required_ids = {s["id"] for s in required}
    for rid in resolved_ids - required_ids:
        r = kb.get(rid)
        if r:
            extra.append({"id": rid, "title": r.get("title", ""),
                          "citation": r.get("citation", "")})

    all_text = "\n\n".join(d.get("text", "") for d in docs)
    platforms = _detect_platforms(all_text)
    total = len(gaps)

    # Flagged items we couldn't map to a KB standard — surface them raw so
    # nothing the platform flagged disappears (Chris matches these by hand).
    unmatched = [
        {"raw": it["raw"], "status": it["status"], "topic": it["topic"],
         "source": it.get("source", "")}
        for it in flagged_items if not it.get("standard_id")
    ]
    matched_flags = [
        {"id": sid, "title": kb.get(sid).get("title", "") if kb.get(sid) else "",
         "citation": it.get("citation", ""), "status": it.get("status", ""),
         "raw": it.get("raw", ""), "source": it.get("source", "")}
        for sid, it in flagged_by_id.items()
    ]

    # Split the score by what Origin can actually act on. Written programs are
    # draftable (Phase 2); References (insurance COIs, TRIR/EMR benchmarks,
    # platform process pages) are context Chris verifies, not documents to draft.
    prog = [g for g in gaps if g["needs_program"]]
    prog_missing = sum(1 for g in prog if g["status"] == "MISSING")
    prog_failing = sum(1 for g in prog if g["status"] == "FAILING")
    prog_present = sum(1 for g in prog if g["status"] == "PRESENT")
    prog_total = len(prog)
    ref_open = sum(1 for g in gaps if not g["needs_program"] and g["status"] != "PRESENT")

    headline = _headline(prog_missing, prog_failing, prog_present, prog_total)

    # Advisory intel — what Origin already knows that's relevant to this trade,
    # these platforms, and the open items. Draws from the REFERENCE brain
    # (prequal-platform how-to, abatement guidance, and anything Origin has been
    # taught), so this panel gets smarter on its own every time the system
    # learns. Retrieval only — it never affects the gap math above or the gate.
    intel: List[dict] = []
    try:
        open_titles = " ".join(g["title"] for g in gaps if g["status"] != "PRESENT")
        q = " ".join(p for p in (industry, " ".join(platforms), open_titles[:400]) if p)
        intel = kb.brain_intel(q, limit=6)
    except Exception:
        intel = []

    return {
        "meta": meta,
        "intel": intel,
        "documents": [{"name": d.get("name", ""),
                       "chars": len(d.get("text", "")),
                       "platforms": _detect_platforms(d.get("text", "")),
                       "is_deficiency_report": looks_like_deficiency_report(d.get("text", ""))}
                      for d in docs],
        "platforms_detected": platforms,
        "deficiency_report": {
            "detected": bool(report_docs),
            "sources": [d.get("name", "") for d in report_docs],
            "flagged_total": len(flagged_items),
            "matched": matched_flags,
            "unmatched": unmatched,
        },
        "summary": {
            "required_total": total,
            "present": counts["PRESENT"],
            "failing": counts["FAILING"],
            "missing": counts["MISSING"],
            # Written-program breakdown — the draftable, actionable core.
            "programs_total": prog_total,
            "programs_present": prog_present,
            "programs_failing": prog_failing,
            "programs_missing": prog_missing,
            "references_open": ref_open,
            "readiness_pct": round(100 * prog_present / prog_total) if prog_total else 0,
            "headline": headline,
        },
        "gaps": gaps,
        "also_covered": extra,
    }


# ── Phase 4: turn a captured LEAD's diagnosis into the exact doc plan ────────
# A contractor who used a free tool told us WHAT is wrong (their trade + the
# issues they flagged, or the OSHA standard they were cited under). This turns
# that into the precise list of documentation that fixes it: every written
# program their trade requires, with the ones tied to their reported problem
# flagged as must-fix so nothing gets missed. Reuses find_gaps so the required
# set stays identical to the internal Gap Finder.

# Rescue-tool issue ids that mean "written programs are missing / rejected /
# don't match scope" — when a lead flags any of these, every required written
# program is a must-fix (that's exactly what they came to us to solve).
_PROGRAM_ISSUE_IDS = {"missing_programs", "scope_mismatch", "outdated_programs",
                      "open_citation"}


def recommend_documents(
    industry: Optional[str] = None,
    state: Optional[str] = None,
    issues: Optional[List[str]] = None,
    citation_program_id: Optional[str] = None,
) -> dict:
    """Produce the exact documentation plan for a lead.

    ``industry``  — the contractor's trade (drives the full required set).
    ``issues``    — rescue-tool issue ids they flagged (marks priorities).
    ``citation_program_id`` — for OSHA-citation leads, the specific program the
                    cited standard requires (always the #1 must-fix).

    Returns ``{industry, have_industry, documents, priority_ids, summary,
    report}`` where each document is
    ``{id, title, citation, category, needs_program, status, priority, reason,
    can_autodraft}`` and priority is one of ``must-fix | required | recommended``.
    """
    issues = issues or []
    prioritize_programs = any(i in _PROGRAM_ISSUE_IDS for i in issues)

    report = None
    have_industry = False
    if industry and industry.strip():
        try:
            report = find_gaps(industry.strip(), state=state, operators=None, docs=[])
            have_industry = True
        except Exception:
            report = None

    documents: List[dict] = []
    seen: set = set()

    def _priority_for(gid: str, needs_program: bool) -> Tuple[str, str]:
        if citation_program_id and gid == citation_program_id:
            return "must-fix", "The written program your OSHA citation requires — build this first."
        if needs_program and prioritize_programs:
            return ("must-fix",
                    "You flagged missing, rejected, or mismatched written programs — "
                    "this one is required for your trade and Origin can draft it now.")
        if needs_program:
            return "required", "A written program your trade is required to have on file."
        return "recommended", "Supporting item to verify (insurance, benchmark, or platform setup)."

    if report:
        for g in report.get("gaps", []):
            needs = bool(g.get("needs_program"))
            prio, reason = _priority_for(g["id"], needs)
            documents.append({
                "id": g["id"],
                "title": g.get("title", ""),
                "citation": g.get("citation", ""),
                "category": g.get("category", ""),
                "needs_program": needs,
                "status": g.get("status", "MISSING"),
                "priority": prio,
                "reason": reason,
                "can_autodraft": bool(g.get("can_autodraft")),
            })
            seen.add(g["id"])

    # Make sure the cited standard's program is present even if we have no trade
    # yet (citation leads don't collect an industry) — never lose the one doc
    # that directly answers their citation.
    if citation_program_id and citation_program_id not in seen:
        rec = kb.get(citation_program_id)
        if rec:
            needs = (rec.get("written_program", "") or "").strip().lower() in ("yes", "conditional")
            documents.append({
                "id": citation_program_id,
                "title": rec.get("title", ""),
                "citation": rec.get("citation", ""),
                "category": rec.get("category", ""),
                "needs_program": needs,
                "status": "MISSING",
                "priority": "must-fix",
                "reason": "The written program your OSHA citation requires — build this first.",
                "can_autodraft": kb.render_program(citation_program_id) is not None,
            })
            seen.add(citation_program_id)

    # Order: must-fix, then required, then recommended; draftable programs before
    # reference items within each tier; alphabetical by title as a tiebreak.
    _rank = {"must-fix": 0, "required": 1, "recommended": 2}
    documents.sort(key=lambda d: (_rank.get(d["priority"], 3),
                                  0 if d["needs_program"] else 1,
                                  d["title"].lower()))

    programs = [d for d in documents if d["needs_program"]]
    summary = {
        "total": len(documents),
        "programs": len(programs),
        "must_fix": sum(1 for d in documents if d["priority"] == "must-fix"),
        "required": sum(1 for d in documents if d["priority"] == "required"),
        "recommended": sum(1 for d in documents if d["priority"] == "recommended"),
    }
    return {
        "industry": (industry or "").strip(),
        "have_industry": have_industry,
        "documents": documents,
        "priority_ids": [d["id"] for d in documents if d["priority"] == "must-fix"],
        "summary": summary,
        "report": report,
    }


# ── Phase 2: auto-draft the missing / failing written programs ──────────────
def _slug(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "-", (s or "").lower()).strip("-")[:60] or "program"


def draft_programs(
    ids: List[str],
    company: Optional[str] = None,
    effective_date: Optional[str] = None,
) -> List[dict]:
    """Render a fillable written program for each standard id.

    Reuses compliance_kb.render_program (which serves the pre-generated template
    or builds one on the fly from the KB record, so a draft can never drift from
    the standard). Pre-fills the company name / effective date when supplied;
    every other {{PLACEHOLDER}} and [[prompt]] is left for Chris to complete.
    """
    company = (company or "").strip()
    effective_date = (effective_date or "").strip()
    out: List[dict] = []
    seen: set = set()
    for eid in ids:
        if not eid or eid in seen:
            continue
        seen.add(eid)
        md = kb.render_program(eid)
        if not md:
            continue
        rec = kb.get(eid) or {}
        if company:
            md = md.replace("{{COMPANY_NAME}}", company)
        if effective_date:
            md = md.replace("{{EFFECTIVE_DATE}}", effective_date)
        title = rec.get("title", eid)
        out.append({
            "id": eid,
            "title": title,
            "citation": rec.get("citation", ""),
            "filename": f"program-{_slug(title)}.md",
            "markdown": md,
        })
    return out


# Brand palette (matches the rest of Origin's compliance deliverables).
_CHARCOAL = "1C1F24"
_ORANGE = "E8551F"
_GREY = "6B7280"

_INLINE_RE = re.compile(r"(\{\{[^}]*\}\}|\[\[[^\]]*\]\]|\*\*[^*]+\*\*|\*[^*]+\*)")


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def _add_runs(paragraph, text: str) -> None:
    """Emit styled runs, highlighting {{PLACEHOLDERS}} (orange bold, still to
    fill) and [[prompts]] (grey italic guidance) so the doc reads as fillable."""
    from docx.shared import RGBColor
    for part in _INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("{{"):
            r = paragraph.add_run(part); r.bold = True
            r.font.color.rgb = RGBColor.from_string(_ORANGE)
        elif part.startswith("[["):
            r = paragraph.add_run(part[2:-2]); r.italic = True
            r.font.color.rgb = RGBColor.from_string(_GREY)
        elif part.startswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def _bottom_border(paragraph, color: str = _ORANGE) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4"); bottom.set(qn("w:color"), color)
    borders.append(bottom)
    p.append(borders)


def _normalize_program_md(md: str) -> str:
    """render_program returns two shapes: a clean on-the-fly markdown build, OR
    a pre-generated static template with YAML frontmatter + an HTML letterhead
    table. Flatten both to plain markdown so the docx builder handles one format:
    drop the frontmatter, and turn any HTML letterhead/rule lines into readable
    (bold) header text."""
    md = md or ""
    # Strip a leading YAML frontmatter block ( --- ... --- ).
    if md.lstrip().startswith("---"):
        start = md.find("---")
        end = md.find("\n---", start + 3)
        if end != -1:
            nl = md.find("\n", end + 1)
            md = md[nl + 1:] if nl != -1 else ""
    out: List[str] = []
    for line in md.splitlines():
        s = line.strip()
        if s.startswith("<"):  # HTML letterhead / rule lines
            t = re.sub(r"(?i)<br\s*/?>", "\n", s)
            t = re.sub(r"(?i)</(td|tr|div|table|p|h\d)>", "\n", t)
            t = re.sub(r"<[^>]+>", "", t)
            for piece in t.split("\n"):
                piece = piece.strip()
                if piece:
                    out.append(f"**{piece}**")
        else:
            out.append(line)
    return "\n".join(out).strip()


def program_docx_bytes(title: str, markdown: str) -> Optional[bytes]:
    """Render one program's markdown (from render_program) into a branded .docx.

    Returns the file bytes, or None if python-docx isn't available (the caller
    falls back to shipping the markdown so a draft is never lost)."""
    if not _docx_available():
        return None
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor

    markdown = _normalize_program_md(markdown)

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Brand strip
    brand = doc.add_paragraph()
    br = brand.add_run("ORIGIN"); br.bold = True; br.font.size = Pt(15)
    br.font.color.rgb = RGBColor.from_string(_CHARCOAL)
    dot = brand.add_run("."); dot.bold = True; dot.font.size = Pt(15)
    dot.font.color.rgb = RGBColor.from_string(_ORANGE)
    _bottom_border(brand)

    for raw in markdown.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if line.startswith("### "):
            h = doc.add_paragraph(); h.space_before = Pt(8)
            r = h.add_run(line[4:]); r.bold = True; r.font.size = Pt(11.5)
            r.font.color.rgb = RGBColor.from_string(_CHARCOAL)
        elif line.startswith("## "):
            h = doc.add_paragraph()
            r = h.add_run(line[3:]); r.bold = True; r.font.size = Pt(13)
            r.font.color.rgb = RGBColor.from_string(_ORANGE)
        elif line.startswith("# "):
            r = doc.add_paragraph().add_run(line[2:]); r.bold = True
            r.font.size = Pt(18); r.font.color.rgb = RGBColor.from_string(_CHARCOAL)
        elif line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("\u2610 ")  # ballot box
            _add_runs(p, line[6:])
        elif line.startswith("- "):
            _add_runs(doc.add_paragraph(style="List Bullet"), line[2:])
        else:
            _add_runs(doc.add_paragraph(), line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _headline(missing: int, failing: int, present: int, total: int) -> str:
    if total == 0:
        return "No required written programs resolved — check the industry/state inputs."
    if missing == 0 and failing == 0:
        return "Every required written program is covered. Ready for review."
    open_items = missing + failing
    return (f"{open_items} of {total} required written programs need work "
            f"({missing} missing, {failing} failing). Origin can draft all of them.")
